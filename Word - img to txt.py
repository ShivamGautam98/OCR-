from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor
import pytesseract
import os
import cv2
import numpy as np
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx2pdf import convert
document = Document()
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract'
unit=input("Enter Unit no. you want to Convert-")
pg=1
for i in range(1,9):
    path1='C:\\Users\\BK GAUTAM\\Desktop\\ML\\Image to Text\\Unit '+str(unit)+'\\u'+str(unit)+'-L'+str(i)+"\\"
    try:
        images = os.listdir(path1)
        print(path1)
        document.add_heading('\n==================LESSON '+str(i)+'=====================\n\n').bold=True
        for im in images:
            document.add_paragraph('\n--------------- '+im+' --------------------\n').bold=True
            p=document.add_paragraph(' ')
            print(im)
            path=path1+im
            img = cv2.imread(r'C:\\Users\\BK GAUTAM\\Desktop\\ML\\Image to Text\\Unit '+str(unit)+'\\u'+str(unit)+'-L'+str(i)+"\\"+im)
            img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            img = img.astype('uint8')
            blur = cv2.GaussianBlur(img,(1,1),0)
            ret3,th3 = cv2.threshold(blur,0,255,cv2.THRESH_BINARY | cv2.THRESH_OTSU)
            image=th3
            #cv2.imshow('C:\\Users\\BK GAUTAM\\Desktop\\Brand_Partner_Dim.png', image)
            p.add_run(pytesseract.image_to_string(image))
            document.add_heading('Pg-'+str(pg)).bold=True
            p.paragraph_format.line_spacing = 1.5
            pg=pg+1
            cv2.waitKey(0)
            cv2.destroyAllWindows()
    except:
            print(path1, "  Not Found")
document.save('C:\\Users\\BK GAUTAM\\Desktop\\ML\\Image to Text\\Unit '+str(unit)+'\\Unit'+str(unit)+'_MEIO_CA.docx')
convert('C:\\Users\\BK GAUTAM\\Desktop\\ML\\Image to Text\\Unit '+str(unit)+'\\Unit'+str(unit)+"_MEIO_CA.docx", "C:\\Users\\BK GAUTAM\\Desktop\\ML\\Image to Text\\Unit "+str(unit)+'\\Unit'+str(unit)+'_MEIO_CA.pdf')

